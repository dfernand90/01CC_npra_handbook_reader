"""
Shared knowledge-base helper for MCP servers.

Discovers files in a server's knowledge/ folder, extracts content from
multiple formats, and exposes them via MCP /resources/list + /resources/read.

SUPPORTED FORMATS:
  Text:   .md .txt .json .yaml .yml
  Office: .pdf .docx .xlsx .csv
  Image:  .png .jpg .jpeg .gif .webp .bmp .tiff .svg

USAGE (2 lines in any server.py):
    from knowledge_helpers import add_knowledge_routes
    add_knowledge_routes(app, server_id="my-server-id")

Images are returned as base64-encoded data URIs.
PDF/DOCX/XLSX are extracted to plain text (requires optional deps).
If a dependency is missing, files of that type are skipped with a warning.
"""
from __future__ import annotations

import base64
import logging
import os
from pathlib import Path
from typing import Any

from fastapi import FastAPI
from pydantic import BaseModel

logger = logging.getLogger("knowledge_helpers")

# ── Optional dependency imports ───────────────────────────────────────

_HAS_PYPDF = False
_HAS_DOCX = False
_HAS_OPENPYXL = False

try:
    import pypdf
    _HAS_PYPDF = True
except ImportError:
    pass

try:
    import docx as _docx_module
    _HAS_DOCX = True
except ImportError:
    pass

try:
    import openpyxl
    _HAS_OPENPYXL = True
except ImportError:
    pass


# ── MIME type mapping ─────────────────────────────────────────────────

_TEXT_EXTENSIONS = {".md", ".txt", ".json", ".yaml", ".yml", ".csv"}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".svg"}
_OFFICE_EXTENSIONS = {".pdf", ".docx", ".xlsx"}

_ALL_SUPPORTED = _TEXT_EXTENSIONS | _IMAGE_EXTENSIONS | _OFFICE_EXTENSIONS

_MIME_MAP = {
    ".md": "text/markdown",
    ".txt": "text/plain",
    ".json": "application/json",
    ".yaml": "text/yaml",
    ".yml": "text/yaml",
    ".csv": "text/csv",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".tiff": "image/tiff",
    ".svg": "image/svg+xml",
}


# ── Content extractors ────────────────────────────────────────────────

def _read_text(filepath: str) -> str:
    """Read a plain text file."""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


def _read_pdf(filepath: str) -> str:
    """Extract text from a PDF file."""
    if not _HAS_PYPDF:
        return "[PDF reading requires 'pypdf'. Install: pip install pypdf]"
    reader = pypdf.PdfReader(filepath)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {i + 1} ---\n{text}")
    if not pages:
        return "[PDF contains no extractable text — may be image-based. Consider OCR.]"
    return "\n\n".join(pages)


def _read_docx(filepath: str) -> str:
    """Extract text from a Word .docx file."""
    if not _HAS_DOCX:
        return "[DOCX reading requires 'python-docx'. Install: pip install python-docx]"
    doc = _docx_module.Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    if not paragraphs:
        return "[DOCX contains no extractable text]"
    return "\n\n".join(paragraphs)


def _read_xlsx(filepath: str) -> str:
    """Extract data from an Excel .xlsx file as text."""
    if not _HAS_OPENPYXL:
        return "[XLSX reading requires 'openpyxl'. Install: pip install openpyxl]"
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"## Sheet: {ws.title}\n" + "\n".join(rows))
    wb.close()
    if not sheets:
        return "[XLSX contains no data]"
    return "\n\n".join(sheets)


def _read_image_base64(filepath: str) -> str:
    """Read an image file and return base64-encoded content."""
    with open(filepath, "rb") as f:
        data = f.read()
    return base64.b64encode(data).decode("ascii")


def _extract_content(filepath: str, ext: str) -> tuple[str, str]:
    """
    Extract content from a file.

    Returns (content, content_mime_type).
    For text-extractable formats, returns (text, "text/...").
    For images, returns (base64_data, "image/...").
    """
    if ext in _TEXT_EXTENSIONS:
        return _read_text(filepath), _MIME_MAP.get(ext, "text/plain")

    if ext == ".pdf":
        return _read_pdf(filepath), "text/plain"

    if ext == ".docx":
        return _read_docx(filepath), "text/plain"

    if ext == ".xlsx":
        return _read_xlsx(filepath), "text/plain"

    if ext in _IMAGE_EXTENSIONS:
        b64 = _read_image_base64(filepath)
        mime = _MIME_MAP.get(ext, "application/octet-stream")
        return b64, mime

    return f"[Unsupported format: {ext}]", "text/plain"


# ── Discovery ─────────────────────────────────────────────────────────

def discover_knowledge(knowledge_dir: str, server_id: str) -> dict[str, dict]:
    """
    Scan a knowledge/ folder and build a URI → metadata map.

    Returns a dict keyed by URI with entries like:
        {
            "uri": "knowledge://server-id/filename-stem",
            "name": "Human Readable Name",
            "mimeType": "text/markdown",
            "description": "Source: filename.md",
            "_filepath": "/full/path/to/file.md",
            "_ext": ".md",
        }
    """
    resources: dict[str, dict] = {}
    if not os.path.isdir(knowledge_dir):
        return resources

    for filename in sorted(os.listdir(knowledge_dir)):
        filepath = os.path.join(knowledge_dir, filename)
        if not os.path.isfile(filepath):
            continue

        ext = Path(filename).suffix.lower()
        if ext not in _ALL_SUPPORTED:
            logger.debug("Skipping unsupported file: %s", filename)
            continue

        stem = Path(filename).stem
        uri = f"knowledge://{server_id}/{stem}"
        name = stem.replace("_", " ").replace("-", " ").title()
        mime = _MIME_MAP.get(ext, "application/octet-stream")

        resources[uri] = {
            "uri": uri,
            "name": name,
            "mimeType": mime,
            "description": f"Source: {filename}",
            "_filepath": filepath,
            "_ext": ext,
        }

    if resources:
        logger.info(
            "Discovered %d knowledge files for %s: %s",
            len(resources), server_id,
            [os.path.basename(r["_filepath"]) for r in resources.values()],
        )

    return resources


# ── FastAPI route factory ─────────────────────────────────────────────

class _ResourceReadRequest(BaseModel):
    uri: str


def add_knowledge_routes(
    app: FastAPI,
    server_id: str,
    knowledge_dir: str | None = None,
) -> dict[str, dict]:
    """
    Add /resources/list and /resources/read MCP endpoints to a FastAPI app.

    Args:
        app: The FastAPI application instance.
        server_id: Unique server identifier (used in knowledge:// URIs).
        knowledge_dir: Path to knowledge/ folder. Defaults to
                       <server.py directory>/knowledge/

    Returns:
        The discovered knowledge map (for reference/logging).

    Usage:
        from knowledge_helpers import add_knowledge_routes
        add_knowledge_routes(app, server_id="ec2-crack-width")
    """
    if knowledge_dir is None:
        # Default: knowledge/ folder next to the importing module
        # Use the app's root_path or fall back to caller's directory
        import inspect
        caller_file = inspect.stack()[1].filename
        knowledge_dir = os.path.join(os.path.dirname(os.path.abspath(caller_file)), "knowledge")

    knowledge_map = discover_knowledge(knowledge_dir, server_id)

    @app.post("/resources/list")
    async def list_resources():
        """List available knowledge documents embedded in this server."""
        resources = [
            {
                "uri": r["uri"],
                "name": r["name"],
                "mimeType": r["mimeType"],
                "description": r.get("description", ""),
            }
            for r in knowledge_map.values()
        ]
        return {"resources": resources}

    @app.post("/resources/read")
    async def read_resource(request: _ResourceReadRequest):
        """Read a knowledge document by its URI."""
        resource = knowledge_map.get(request.uri)
        if not resource:
            return {"error": f"Unknown resource URI: {request.uri}"}
        try:
            content, content_mime = _extract_content(
                resource["_filepath"], resource["_ext"]
            )
            return {
                "uri": request.uri,
                "name": resource["name"],
                "mimeType": content_mime,
                "content": content,
            }
        except Exception as e:
            logger.error("Failed to read %s: %s", request.uri, e)
            return {"error": f"Failed to read resource: {e}"}

    return knowledge_map
